import streamlit as st
import os
import requests
import json
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import pandas as pd

# ==========================================
# CONFIG & INITIALIZATION
# ==========================================
TEMPLATE_DIR = "hr_templates"
OUTPUT_DIR = "generated_docs"

# สร้างโฟลเดอร์สำหรับเก็บข้อมูลหากยังไม่มี
for folder in [TEMPLATE_DIR, OUTPUT_DIR]:
    if not os.path.exists(folder):
        os.makedirs(folder)

st.set_page_config(page_title="HR AI Document Automation", layout="wide")
st.title("📄 ระบบเว็บแอปพลิเคชันจัดการเอกสาร HR อัจฉริยะ")
st.subheader("ระบบอัตโนมัติ: เพิ่ม/ลบเทมเพลต และรองรับไฟล์นำเข้าหลากรูปแบบด้วย Open-Source AI")

# ==========================================
# FUNCTIONS (ระบบเบื้องหลัง)
# ==========================================

# 1. ฟังก์ชันอ่านข้อมูลจากไฟล์นำเข้า (Input) รองรับ ภาพ, PDF, Excel
def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    
    if ext in ['.jpg', '.jpeg', '.png']:
        # อ่านข้อความจากรูปภาพ (OCR)
        text = pytesseract.image_to_string(Image.open(file_path), lang='tha+eng')
    elif ext == '.pdf':
        # แปลง PDF เป็นภาพแล้วทำ OCR
        pages = convert_from_path(file_path)
        for page in pages:
            text += pytesseract.image_to_string(page, lang='tha+eng') + "\n"
    elif ext == '.xlsx':
        # อ่านไฟล์ Excel ดิจิทัลโดยตรง
        try:
            df = pd.read_excel(file_path)
            text = "ข้อมูลจากไฟล์ Excel:\n"
            text += df.to_string(index=False)
        except Exception as e:
            st.error(f"ไม่สามารถอ่านไฟล์ Excel ได้: {e}")
            
    return text

# 2. ฟังก์ชันส่งข้อความให้ AI (Ollama - โมเดลภาษาไทยฟรี ไม่มีลิมิต)
def extract_data_with_ai(raw_text, keys_needed):
    # เชื่อมต่อกับ Ollama ที่เปิดรันไว้ในเครื่องคอมพิวเตอร์ของคุณ
    url = "http://localhost:11434/api/generate" 
    
    prompt = f"""
    คุณคือ AI ผู้เชี่ยวชาญงานฝ่ายบุคคล (HR) จงอ่านข้อความที่กำหนดให้ต่อไปนี้ 
    แล้วดึงข้อมูลสำคัญออกมาจับคู่กับหัวข้อ (Keys) ที่ระบุ โดยตอบกลับในรูปแบบ JSON object เท่านั้น ห้ามมีคำอธิบายอื่น

    หัวข้อที่ต้องการจับคู่ (Keys): {keys_needed}

    ข้อความเอกสารนำเข้า:
    \"\"\"
    {raw_text}
    \"\"\"

    ตอบกลับเป็น JSON เท่านั้น:
    """
    
    data = {
        "model": "typhoon-m1", # สามารถเปลี่ยนเป็น llama3.1 หรือโมเดลอื่นที่โหลดไว้ได้
        "prompt": prompt,
        "stream": False,
        "format": "json"
    }
    
    try:
        response = requests.post(url, json=data)
        result = response.json()
        return json.loads(result['response'])
    except Exception as e:
        st.error(f"ระบบไม่สามารถเชื่อมต่อกับ AI Model ได้ (ตรวจสอบการเปิด Ollama): {e}")
        return None

# 3. ฟังก์ชันนำข้อมูล JSON ไปเขียนแทนที่ตัวแปรในไฟล์เทมเพลต Word (.docx)
def fill_template(template_path, data, output_path):
    doc = Document(template_path)
    
    # เปลี่ยนข้อความในย่อหน้าปกติ
    for p in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, str(value))
                
    # เปลี่ยนข้อความที่อยู่ในตารางภายในไฟล์ Word
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(value))
                            
    doc.save(output_path)

# ==========================================
# UI SIDEBAR: ระบบจัดการเทมเพลต (ข้อ 2)
# ==========================================
with st.sidebar:
    st.header("⚙️ คลังจัดการเทมเพลตเอกสาร")
    
    # ฟังก์ชันอัปโหลดเทมเพลต (.docx สำหรับนำไปกรอกข้อความอัตโนมัติ)
    uploaded_tpl = st.file_uploader("📥 อัปโหลดเทมเพลตใหม่ (.docx เท่านั้น)", type=["docx"])
    if uploaded_tpl is not None:
        tpl_path = os.path.join(TEMPLATE_DIR, uploaded_tpl.name)
        with open(tpl_path, "wb") as f:
            f.write(uploaded_tpl.getbuffer())
        st.success(f"บันทึกเทมเพลต '{uploaded_tpl.name}' เข้าสู่คลังแล้ว!")
        st.rerun()

    st.divider()
    
    # ฟังก์ชันลบเทมเพลตออกจากระบบ
    st.subheader("🗑️ ลบเทมเพลตที่ไม่ได้ใช้งาน")
    all_tpls = os.listdir(TEMPLATE_DIR)
    if all_tpls:
        tpl_to_delete = st.selectbox("เลือกไฟล์ที่ต้องการลบ:", all_tpls, key="delete_box")
        if st.button("🔴 ยืนยันการลบไฟล์ถาวร"):
            os.remove(os.path.join(TEMPLATE_DIR, tpl_to_delete))
            st.success(f"ลบไฟล์ {tpl_to_delete} เรียบร้อยแล้ว")
            st.rerun()
    else:
        st.caption("ยังไม่มีไฟล์เทมเพลตในคลังข้อมูล")

# ==========================================
# MAIN UI: ส่วนการทำงานหลักของ HR (ข้อ 1, 3, 4, 5)
# ==========================================
available_templates = os.listdir(TEMPLATE_DIR)

if not available_templates:
    st.info("💡 เริ่มต้นใช้งานโดยการอัปโหลดไฟล์เทมเพลตเอกสารของบริษัท (.docx) ที่เมนูด้านซ้ายมือก่อนครับ")
else:
    # ขั้นตอนที่ 1: การเลือกใช้เทมเพลต (ข้อ 3)
    st.header("Step 1: เลือกเทมเพลตจากระบบ")
    selected_template = st.selectbox("เลือกเทมเพลตที่คุณต้องการนำข้อมูลไปกรอก:", available_templates, key="main_select")
    
    # กำหนดตัวแปรเพื่อให้สอดคล้องกับเครื่องหมาย {{ตัวแปร}} ในไฟล์ Word
    st.markdown("*ระบุชื่อตัวแปรที่ต้องการให้ AI เข้าไปแทนที่ในเทมเพลต (ใช้จุลภาค `,` คั่น)*")
    fields_input = st.text_input("ชื่อฟิลด์ข้อมูลที่ต้องการ:", "Name, ID_Number, Address, Position, Salary")
    fields_list = [f.strip() for f in fields_input.split(",")]

    st.divider()

    # ขั้นตอนที่ 2: อัปโหลดไฟล์ข้อมูลดิบนำเข้า (ข้อ 4)
    st.header("Step 2: อัปโหลดไฟล์เอกสาร Input นำเข้า")
    input_file = st.file_uploader(
        "ลากหรือเลือกไฟล์ข้อมูลดิบ (รองรับไฟล์ภาพ .png, .jpg, .jpeg | ไฟล์เอกสาร .pdf | ไฟล์ตาราง .xlsx)", 
        type=["png", "jpg", "jpeg", "pdf", "xlsx"]
    )

    if input_file is not None:
        # บันทึกไฟล์ข้อมูลดิบชั่วคราวลงในโฟลเดอร์ผลลัพธ์
        temp_input_path = os.path.join(OUTPUT_DIR, input_file.name)
        with open(temp_input_path, "wb") as f:
            f.write(input_file.getbuffer())
            
        if st.button("🚀 สั่ง AI ดึงข้อมูลและกรอกเอกสารแทนคน"):
            # แสดงสถานะการทำงานเป็นขั้นตอนอย่างชัดเจน (ลดงานคีย์ข้อความ)
            with st.status("🤖 AI กำลังจัดการเอกสารแทนคุณ...", expanded=True) as status:
                
                status.write("⏳ 1. กำลังสแกนและแกะข้อความจากไฟล์นำเข้าของคุณ...")
                raw_text = extract_text_from_file(temp_input_path)
                
                status.write("🧠 2. กำลังส่งข้อมูลให้ Open-Source AI ทำการวิเคราะห์และจับคู่ลงโครงร่าง...")
                extracted_json = extract_data_with_ai(raw_text, fields_list)
                
                if extracted_json:
                    status.write("📝 3. ผลลัพธ์ข้อมูลที่ AI ตรวจจับและดึงออกมาได้สำเร็จ:")
                    status.json(extracted_json)
                    
                    # ตั้งชื่อไฟล์และกำหนดโฟลเดอร์สำหรับส่งออกไฟล์ใหม่
                    output_filename = f"Filled_{selected_template}"
                    output_path = os.path.join(OUTPUT_DIR, output_filename)
                    template_path = os.path.join(TEMPLATE_DIR, selected_template)
                    
                    status.write("💾 4. กำลังรวบรวมข้อมูลเขียนลงในไฟล์เทมเพลตอัตโนมัติ...")
                    fill_template(template_path, extracted_json, output_path)
                    
                    status.update(label="🎉 ดำเนินการเสร็จสิ้น! เอกสารของคุณพร้อมใช้งานแล้ว", state="complete", expanded=True)
                    
                    # แสดงกล่องแจ้งเตือนความสำเร็จและสร้างปุ่มให้ดาวน์โหลดไฟล์ทันที
                    st.success(f"ระบบจัดทำเอกสารเสร็จสมบูรณ์!")
                    with open(output_path, "rb") as file:
                        st.download_button(
                            label="📥 ดาวน์โหลดไฟล์เอกสารสำเร็จรูป (.docx)",
                            data=file,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    status.update(label="❌ เกิดข้อผิดพลาดในขั้นตอนประมวลผล", state="error")
                    st.error("AI ไม่สามารถดึงข้อมูลออกมาได้ กรุณาตรวจสอบเนื้อหาของไฟล์หรือฟิลด์ข้อมูลที่ระบุ")
